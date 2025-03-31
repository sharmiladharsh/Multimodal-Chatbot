import os
import io
import cv2
import faiss
import torch
import numpy as np
import pandas as pd
from PIL import Image
import streamlit as st
from docx import Document
from transformers import pipeline, BlipProcessor, BlipForQuestionAnswering
from sentence_transformers import SentenceTransformer

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Load models
st.spinner("Loading SentenceTransformer...")
embedder = SentenceTransformer("all-MiniLM-L6-v2")
qa_model = pipeline("question-answering", model="bert-large-uncased-whole-word-masking-finetuned-squad")

def extract_text(file, file_type):
    file_bytes = file.read()

    if file_type == "pdf":
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "".join([page.get_text() for page in doc]), None

    elif file_type == "docx":
        doc = Document(io.BytesIO(file_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        dataframes = []
        for table in doc.tables:
            keys = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                if len(values) == len(keys):
                    dataframes.append(dict(zip(keys, values)))
        df = pd.DataFrame(dataframes) if dataframes else None
        return full_text, df

    elif file_type == "csv":
        df = pd.read_csv(io.BytesIO(file_bytes))
        return df.to_string(index=False), df

    return "", None

def chunk_text(text, max_tokens=300, overlap=50):
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i+max_tokens]
        chunks.append(" ".join(chunk))
        i += max_tokens - overlap
    return chunks[:30]

def create_index(chunks):
    embeddings = embedder.encode(chunks)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(np.array(embeddings))
    return index, chunks

def get_top_chunks(question, chunks, index, k=3):
    q_vec = embedder.encode([question])
    _, I = index.search(np.array(q_vec), k)
    return [chunks[i] for i in I[0]]

def generate_answer(context, question):
    result = qa_model(question=question, context=context)
    return result['answer'].strip()

def handle_structured_question(df, question):
    q = question.lower()
    df.columns = [col.strip().lower() for col in df.columns]
    numeric_cols = df.select_dtypes(include='number').columns

    def match_column_from_question():
        for col in numeric_cols:
            if col in q:
                return col
        return None

    if "lowest" in q and "average" in q:
        averages = {col: df[col].mean() for col in numeric_cols}
        if averages:
            lowest = min(averages, key=averages.get)
            return f"{lowest} has the lowest average: {averages[lowest]:.2f}"

    if "highest" in q and "score" in q:
        matched_col = match_column_from_question()
        if matched_col:
            row = df[df[matched_col] == df[matched_col].max()]
            return f"{row.iloc[0][0]} scored the highest in {matched_col} with {row.iloc[0][matched_col]}"

    if "average" in q:
        matched_col = match_column_from_question()
        if matched_col:
            return f"The average {matched_col} score is {df[matched_col].mean():.2f}"

    if "how many" in q:
        for col in df.columns:
            if col in q:
                value_counts = df[col].value_counts()
                for val in value_counts.index:
                    if val.lower() in q:
                        return f"{value_counts[val]} students got {val}"
                return str(value_counts.to_dict())

    return None

def blip_answer(image_path, question):
    processor = BlipProcessor.from_pretrained("Salesforce/blip-vqa-base")
    model = BlipForQuestionAnswering.from_pretrained("Salesforce/blip-vqa-base").to("cpu")
    image = Image.open(image_path).convert("RGB")
    inputs = processor(image, question, return_tensors="pt")
    out = model.generate(**inputs)
    return processor.decode(out[0], skip_special_tokens=True)

def extract_frame(video_path, timestamp=1.0):
    output_path = os.path.join(UPLOAD_FOLDER, "frame.jpg")
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)
    frame_number = int(fps * timestamp)
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
    success, frame = cap.read()
    if success:
        cv2.imwrite(output_path, frame)
        return output_path
    return None

# ======================= STREAMLIT UI =======================

st.title("Multimodal Q&A Chatbot")
uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, CSV, Image or Video)", type=["pdf", "docx", "csv", "jpg", "jpeg", "png", "mp4", "avi", "mov"])
question = st.text_input("Ask a question about the uploaded file:")

if uploaded_file and question:
    if st.button("Ask"):
        with st.spinner("Processing..."):
            ext = os.path.splitext(uploaded_file.name)[1].lower()

            temp_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.read())

            if ext in [".pdf", ".docx", ".csv"]:
                raw_text, df = extract_text(open(temp_path, "rb"), ext[1:])
                chunks = chunk_text(raw_text)
                index, chunks = create_index(chunks)

                if df is not None:
                    answer = handle_structured_question(df, question) or generate_answer("\n".join(get_top_chunks(question, chunks, index)), question)
                else:
                    context = "\n".join(get_top_chunks(question, chunks, index))
                    answer = generate_answer(context, question)

            elif ext in [".jpg", ".jpeg", ".png"]:
                answer = blip_answer(temp_path, question)

            elif ext in [".mp4", ".avi", ".mov"]:
                frame_path = extract_frame(temp_path)
                if frame_path:
                    answer = blip_answer(frame_path, question)
                else:
                    answer = "Could not extract frame from video."

            else:
                answer = "Unsupported file type."

            st.success("Answer:")
            st.write(answer)